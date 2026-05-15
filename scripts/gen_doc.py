from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

NEGRO     = RGBColor(0x1A, 0x1A, 0x1A)
GRIS_DARK = RGBColor(0x44, 0x44, 0x44)
GRIS_MID  = RGBColor(0x88, 0x88, 0x88)
ROJO      = RGBColor(0xD7, 0x26, 0x3D)
NARANJA   = RGBColor(0xE8, 0x83, 0x1A)
VERDE     = RGBColor(0x2A, 0x9D, 0x5C)
AZUL      = RGBColor(0x1A, 0x6F, 0xC4)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)

BG_ROJO    = "FFE5E8"
BG_NARANJA = "FFF3E0"
BG_VERDE   = "E8F5EE"
BG_AZUL    = "E3F0FB"
BG_HEADER  = "1A1A1A"


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def run(para, text, bold=False, italic=False, size=11, color=NEGRO, font='Calibri'):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


# ── Portada ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
run(p, 'DATA CLEANING', bold=True, size=30, color=NEGRO)

p2 = doc.add_paragraph()
run(p2, 'Alquileres CABA 2026 — ZonaProp', size=13, color=GRIS_MID)

p3 = doc.add_paragraph()
run(p3, 'Dataset raw: 24.527 filas, 67 columnas', size=11, color=GRIS_MID)
run(p3, '   →   Dataset limpio estimado: ~15.000 filas, ~30 columnas', size=11, color=VERDE)

doc.add_paragraph()

# ── Sección 1 — columnas ──────────────────────────────────────────────────────
s1 = doc.add_paragraph()
run(s1, '01 / Decisiones por columna', bold=True, size=15, color=NEGRO)

ley = doc.add_paragraph()
run(ley, 'Eliminar', bold=True, size=9, color=ROJO)
run(ley, '    Zona gris', bold=True, size=9, color=NARANJA)
run(ley, '    Conservar', bold=True, size=9, color=VERDE)
doc.add_paragraph()

COLUMNAS = [
    # (nombre, decision, razon, bg)
    ('listing_id',        'Eliminar',  'Metadata — no aporta información predictiva',                     BG_ROJO),
    ('source',            'Eliminar',  'Constante: siempre "zonaprop"',                                   BG_ROJO),
    ('detail_url',        'Eliminar',  'Metadata — URL del aviso',                                        BG_ROJO),
    ('scraped_at',        'Eliminar',  'Metadata — timestamp del scraping',                               BG_ROJO),
    ('codigo_aviso',      'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('raw_data',          'Eliminar',  'JSON crudo — redundante con las otras columnas',                  BG_ROJO),
    ('address',           'Eliminar',  'Texto libre — reemplazado por neighborhood',                      BG_ROJO),
    ('neighborhood',      'Conservar', 'Feature clave del modelo. Eliminar filas sin valor (27% nulos)',  BG_VERDE),
    ('city',              'Eliminar',  'Constante: siempre "Buenos Aires"',                               BG_ROJO),
    ('latitude',          'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('longitude',         'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('price_usd',         'Conservar', 'Target del modelo. 0% nulos. Incluye conversion ARS a USD',       BG_VERDE),
    ('price_ars',         'Eliminar',  'Redundante: conversion ya esta en price_usd',                     BG_ROJO),
    ('currency',          'Eliminar',  'Redundante una vez consolidado price_usd',                        BG_ROJO),
    ('price_raw',         'Eliminar',  'Texto original del precio — redundante',                          BG_ROJO),
    ('expenses_ars',      'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('expenses_raw',      'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('exchange_rate',     'Eliminar',  'Constante ($1.280). Redundante post-consolidacion',               BG_ROJO),
    ('usd_equivalent',    'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('property_type',     'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('operation_type',    'Eliminar',  'Constante: siempre "alquiler"',                                   BG_ROJO),
    ('ambientes',         'Conservar', '2.7% nulos — eliminar filas. Feature importante',                 BG_VERDE),
    ('dormitorios',       'Zona gris', '28% nulos — imputar con ambientes - 1 (valido en 83%)',           BG_NARANJA),
    ('banos',             'Conservar', '3.8% nulos — eliminar filas',                                     BG_VERDE),
    ('toilettes',         'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('m2_total',          'Conservar', '1.2% nulos — eliminar filas. Feature importante',                 BG_VERDE),
    ('m2_covered',        'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('m2_uncovered',      'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('m2_semicovered',    'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('antiguedad',        'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('piso',              'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('pisos_edificio',    'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('orientacion',       'Eliminar',  '95.8% nulos',                                                     BG_ROJO),
    ('disposicion',       'Eliminar',  '55.9% nulos',                                                     BG_ROJO),
    ('luminosidad',       'Eliminar',  '55.4% nulos. Si se conserva: convertir a numerica (0/1/2)',       BG_ROJO),
    ('estado',            'Eliminar',  '80.8% nulos. 6 categorias ordinales pero datos insuficientes',    BG_ROJO),
    ('tipo_calefaccion',  'Eliminar',  '88.8% nulos',                                                     BG_ROJO),
    ('cochera (bool)',    'Eliminar',  'Reemplazada por cochera_cantidad — mas informacion',               BG_ROJO),
    ('cochera_cantidad',  'Conservar', '0% nulos. Entero: 0 = sin cochera, 1+ = cantidad',               BG_VERDE),
    ('baulera',           'Conservar', '0% nulos',                                                        BG_VERDE),
    ('balcon',            'Conservar', '0% nulos. 57% True — buena varianza',                            BG_VERDE),
    ('terraza',           'Conservar', '0% nulos. 17.9% True',                                           BG_VERDE),
    ('patio',             'Conservar', '0% nulos',                                                        BG_VERDE),
    ('jardin',            'Conservar', '0% nulos',                                                        BG_VERDE),
    ('parrilla',          'Conservar', '0% nulos. 20.3% True',                                           BG_VERDE),
    ('pileta',            'Conservar', '0% nulos. 25.2% True',                                           BG_VERDE),
    ('gimnasio',          'Conservar', '0% nulos. 17.6% True',                                           BG_VERDE),
    ('laundry',           'Conservar', '0% nulos',                                                        BG_VERDE),
    ('sum',               'Conservar', '0% nulos. 31.9% True',                                           BG_VERDE),
    ('seguridad_24hs',    'Conservar', '0% nulos',                                                        BG_VERDE),
    ('ascensor',          'Conservar', '0% nulos',                                                        BG_VERDE),
    ('solarium',          'Conservar', '0% nulos',                                                        BG_VERDE),
    ('cowork',            'Conservar', '0% nulos',                                                        BG_VERDE),
    ('bicipuerto',        'Zona gris', '0% nulos pero solo 1.3% True — casi cero varianza',               BG_NARANJA),
    ('aire_acondicionado','Conservar', '0% nulos. 33.7% True',                                           BG_VERDE),
    ('calefaccion',       'Conservar', '0% nulos',                                                        BG_VERDE),
    ('amenities_raw',     'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('acepta_mascotas',   'Conservar', '0% nulos',                                                        BG_VERDE),
    ('amoblado',          'Conservar', '0% nulos',                                                        BG_VERDE),
    ('apto_profesional',  'Conservar', '0% nulos',                                                        BG_VERDE),
    ('apto_credito',      'Conservar', '0% nulos',                                                        BG_VERDE),
    ('publisher_type',    'Eliminar',  'Un solo valor real ("inmobiliaria") — cero varianza efectiva',    BG_ROJO),
    ('publisher_name',    'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('publisher_phone',   'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('title',             'Eliminar',  'Texto libre — no usable en regresion lineal',                     BG_ROJO),
    ('description',       'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('images',            'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('video_url',         'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('tour_360_url',      'Eliminar',  '100% nulos',                                                      BG_ROJO),
    ('fecha_publicacion', 'Eliminar',  '100% nulos',                                                      BG_ROJO),
]

DCOL = {'Eliminar': ROJO, 'Conservar': VERDE, 'Zona gris': NARANJA}

t1 = doc.add_table(rows=1, cols=3)
t1.style = 'Table Grid'
hdr = t1.rows[0].cells
for cell in hdr:
    set_cell_bg(cell, BG_HEADER)
for cell, txt in zip(hdr, ['Columna', 'Decision', 'Razon']):
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLANCO; r.font.name = 'Calibri'

for col_name, decision, razon, bg in COLUMNAS:
    row = t1.add_row()
    for cell in row.cells:
        set_cell_bg(cell, bg)
    r0 = row.cells[0].paragraphs[0].add_run(col_name)
    r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = NEGRO; r0.font.name = 'Courier New'
    r1 = row.cells[1].paragraphs[0].add_run(decision)
    r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = DCOL.get(decision, NEGRO); r1.font.name = 'Calibri'
    r2 = row.cells[2].paragraphs[0].add_run(razon)
    r2.font.size = Pt(9); r2.font.color.rgb = GRIS_DARK; r2.font.name = 'Calibri'

doc.add_paragraph()

# ── Sección 2 — filas ─────────────────────────────────────────────────────────
s2 = doc.add_paragraph()
run(s2, '02 / Tratamiento de filas', bold=True, size=15, color=NEGRO)
doc.add_paragraph()

FILAS = [
    ('price_usd nulo',      '134 filas',    'Eliminar', 'No se puede imputar el target'),
    ('price_usd < $300',    '~87 filas',    'Eliminar', 'Errores de carga o valores en otra unidad'),
    ('price_usd > $5.000',  '~558 filas',   'Eliminar', 'Outliers — distorsionan el modelo'),
    ('neighborhood nulo',   '~6.700 filas', 'Eliminar', 'Feature critica — no se puede imputar'),
    ('m2_total nulo',       '~302 filas',   'Eliminar', 'Feature importante, muestra chica para imputar'),
    ('m2_total < 15 m2',    'Variable',     'Eliminar', 'Valores imposibles o errores'),
    ('m2_total > 500 m2',   'Variable',     'Eliminar', 'Outliers que distorsionan la regresion'),
    ('ambientes nulo',      '~665 filas',   'Eliminar', 'Feature importante'),
    ('banos nulo',          '~935 filas',   'Eliminar', 'Feature importante'),
    ('dormitorios nulo',    '~6.875 filas', 'Imputar',  'dormitorios = ambientes - 1 (valido en 83% de los casos)'),
]

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
hdr2 = t2.rows[0].cells
for cell in hdr2:
    set_cell_bg(cell, BG_HEADER)
for cell, txt in zip(hdr2, ['Condicion', 'Filas', 'Accion', 'Justificacion']):
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLANCO; r.font.name = 'Calibri'

for cond, n, accion, just in FILAS:
    row = t2.add_row()
    bg = BG_ROJO if accion == 'Eliminar' else BG_NARANJA
    for cell in row.cells:
        set_cell_bg(cell, bg)
    color_accion = ROJO if accion == 'Eliminar' else NARANJA
    for cell, txt, bold, col in zip(
        row.cells, [cond, n, accion, just],
        [True, False, True, False],
        [NEGRO, GRIS_DARK, color_accion, GRIS_DARK]
    ):
        r = cell.paragraphs[0].add_run(txt)
        r.bold = bold; r.font.size = Pt(9); r.font.color.rgb = col; r.font.name = 'Calibri'

doc.add_paragraph()

# ── Sección 3 — feature engineering ──────────────────────────────────────────
s3 = doc.add_paragraph()
run(s3, '03 / Features nuevas (Feature Engineering)', bold=True, size=15, color=NEGRO)
doc.add_paragraph()

FEATURES = [
    ('precio_por_m2',           'price_usd / m2_total',                                              'Metrica estandar del mercado inmobiliario'),
    ('precio_por_ambiente',     'price_usd / ambientes',                                             'Alternativa cuando m2 no esta disponible'),
    ('m2_por_ambiente',         'm2_total / ambientes',                                              'Indica si los ambientes son amplios o chicos'),
    ('m2_por_dormitorio',       'm2_total / dormitorios',                                            'Similar pero enfocado en habitaciones'),
    ('ratio_banos_ambientes',   'banos / ambientes',                                                 'Proxy de calidad del departamento'),
    ('es_monoambiente',         'ambientes == 1',                                                    'Segmento con logica de precio propia'),
    ('amenity_score_edificio',  'sum(pileta, gimnasio, sum, parrilla, seguridad_24hs, ascensor...)', 'Score de servicios del edificio'),
    ('amenity_score_depto',     'sum(balcon, terraza, patio, jardin, baulera, aire_ac...)',          'Score de amenities propios del depto'),
    ('amenity_score_total',     'suma de todos los booleanos',                                       'Score general de equipamiento'),
    ('tiene_espacio_exterior',  'balcon OR terraza OR patio OR jardin',                              'Muy valorado post-pandemia'),
    ('tiene_clima_completo',    'aire_acondicionado AND calefaccion',                                'Confort termico en ambas estaciones'),
    ('es_llave_en_mano',        'amoblado AND aire_acondicionado AND calefaccion',                   'Entras y no gastas nada extra'),
    ('audiencia_amplia',        'apto_profesional OR apto_credito OR acepta_mascotas',               'Cuantas restricciones levanta el propietario'),
    ('es_edificio_premium',     'pileta AND (gimnasio OR sum) AND seguridad_24hs',                   'Edificio con servicios completos'),
]

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
hdr3 = t3.rows[0].cells
for cell in hdr3:
    set_cell_bg(cell, BG_HEADER)
for cell, txt in zip(hdr3, ['Feature nueva', 'Formula', 'Descripcion']):
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLANCO; r.font.name = 'Calibri'

for feat, formula, desc in FEATURES:
    row = t3.add_row()
    for cell in row.cells:
        set_cell_bg(cell, BG_AZUL)
    r0 = row.cells[0].paragraphs[0].add_run(feat)
    r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = AZUL; r0.font.name = 'Courier New'
    r1 = row.cells[1].paragraphs[0].add_run(formula)
    r1.font.size = Pt(8.5); r1.font.color.rgb = GRIS_DARK; r1.font.name = 'Courier New'
    r2 = row.cells[2].paragraphs[0].add_run(desc)
    r2.font.size = Pt(9); r2.font.color.rgb = GRIS_DARK; r2.font.name = 'Calibri'

doc.add_paragraph()
foot = doc.add_paragraph()
run(foot, 'Generado: 2026-04-29  |  Dataset: alquileres_20260302_100514.csv  |  v1.0', size=8, color=GRIS_MID, italic=True)

out = r'e:\ORDENADOR\trabajos\proyectos_personales\inmobiliario_real\docs\data_cleaning.docx'
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f'Guardado: {out}')
